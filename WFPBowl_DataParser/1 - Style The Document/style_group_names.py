# RUN:
# find ingredient groups
# get e.g. location and split following ingredients into a list
# get a list of all the group name:group items from the json file
#     for each group name in the list, if the group name is in the name the currently checked ingredient group, replace the currently checked group with the group name
#     for each group item in the list, if the group item is in the currently checked ingredient group item list, replace the currently checked group items with the item


import docx
import os
import json
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement


class Group:
    def __init__(self, name, items):
        self.name = name
        self.items = items

    def __repr__(self):
        return (f"\n\nGroup: {self.name}"
                f"\nitems: {self.items}\n")

def get_paragraphs(document):
    return [paragraph for paragraph in document.paragraphs]

def get_groups_from_json(json_file):
    f = open(json_file)
    data = json.load(f)
    return [Group(i["name"], i["items"]) for i in data]

def create_group(paragraph):
    text = paragraph.text
    group_name = text[0:text.index(", e.g.")].lower().lstrip("*").strip() 
    group_items = text[text.index(", e.g.") + len(", e.g., "):].lower().strip().split(", ")
    return Group(group_name, group_items)

def get_style(paragraph):
    if "Group" in paragraph.style.name:
        return paragraph.style.name.strip("Group")
    else:
        return paragraph.style.name

def run(path, new_file):
    document = docx.Document(path)
    paragraphs = get_paragraphs(document)
    for paragraph in document.paragraphs:
        if ", e.g." in paragraph.text:
            paragraph_group = create_group(paragraph)
            group_name = paragraph_group.name
            paragraph.text = group_name
            for item in paragraph_group.items:
                paragraph.insert_paragraph_before(f"[{group_name}] {group_name}, {item}", document.styles[get_style(paragraph)])
    document.save(new_file)


if __name__ == '__main__':
    path = "../1 - document_styler/StyledIngredients.docx"
    print(os.path.abspath(path))
    doc = docx.Document(path)

    save_file = "UngroupedIngredients.docx"

    run(path, save_file)