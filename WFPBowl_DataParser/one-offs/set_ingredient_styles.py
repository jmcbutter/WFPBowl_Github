import docx
import os
from docx.enum.text import WD_BREAK

acceptable_styles_list = \
[\
'IngredientName',
'Group',
'Cuisine',
'Category',
'Dish',
'Match1',
'Match2',
'Match3',
'Match4',
'MatchGroup1',
'MatchGroup2',
'MatchGroup3',
'MatchGroup4',
]


#Check Paragraphs In Document To See If They Need To Be Updated
def is_empty(paragraph):
    return not paragraph.text.strip()

def is_ingredient(paragraph):
    if not paragraph.style.name == "IngredientName":
        try:
            return paragraph.runs[0].font.color.rgb == docx.shared.RGBColor(102, 139, 36)
        except:
            return

def write_paragraphs_to_file(path):
    document = docx.Document(path)
    paragraphs = [paragraph for paragraph in document.paragraphs]
    styles = document.styles
    for i in range(0, len(paragraphs)):
        if is_ingredient(paragraphs[i]):
            # try:
                document.paragraphs[i].style = styles["IngredientName"]
                document.paragraphs[i-1].add_run()
                document.paragraphs[i-1].runs[-1].add_break(WD_BREAK.PAGE)
                print(i)
            # except:
                # pass
        else:
            try:
                if paragraphs[i].runs[0].font.color.rgb == docx.shared.RGBColor(0, 99, 137):
                    for run in paragraphs[i].runs:
                        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            except:
                pass
    document.save(path)








path = "../../../IngredientPairsTest.docx"
print(os.path.abspath(path))

doc = docx.Document(path)
write_paragraphs_to_file(path)