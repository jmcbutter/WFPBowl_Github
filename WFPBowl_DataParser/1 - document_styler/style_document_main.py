import docx
import os

acceptable_styles_list = \
[\
'IngredientName',
'Group',
'Cuisine',
'Category',
'Dish',
'Match1',
'Match2',
'Match3',
'Match4',
'MatchGroup1',
'MatchGroup2',
'MatchGroup3',
'MatchGroup4',
]

#Check Paragraphs In Document To See If They Need To Be Updated
def is_empty(paragraph):
    return not paragraph.text.strip()

def is_ingredient(paragraph):
    return paragraph.style.name == "IngredientName"

def is_heading(paragraph):
    return "heading" in paragraph.style.name.lower() 

def requires_update(paragraph):
    if not is_empty(paragraph) and not is_ingredient(paragraph) and not is_heading(paragraph):
        if paragraph.style.name not in acceptable_styles_list:
            return True
    return False


#Determine the Match Strength based on text boldness, capitalization, italicization.
#Only checks the first run of the paragraph in case later words/letters aren't the same style
def has_asterisk(paragraph):
    return "*" in paragraph.runs[0].text

def is_bold(paragraph):
    return paragraph.runs[0].bold

def is_all_caps(paragraph):
    return paragraph.runs[0].text[0] == paragraph.runs[0].text[0].upper()

def get_match_strength(paragraph):
    if has_asterisk(paragraph):
        return 1
    else:
        if is_all_caps(paragraph):
            return 2
        elif is_bold(paragraph):
            return 3
        else:
            return 4
    raise Exception(f"There's a problem with one of the ingredients... Paragraph Text:\n\t {paragraph.text}")


#Determines Which Style To Apply Based On Match Strength or Italicization
def is_italicized(paragraph):
    return paragraph.runs[0].italic

def is_cuisine(paragraph):
    return "cuisine" in paragraph.text.lower()

def determine_style(paragraph):
    if is_italicized(paragraph):
        if is_cuisine(paragraph):
            return 'Cuisine'
        else:
            return 'Dish'
    elif ", e.g." in paragraph.text:
        return "GroupMatch" + str(get_match_strength(paragraph))
    else:
        return "Match" + str(get_match_strength(paragraph))

#Set color based on style
def determine_font_RGB_color(style):
    if style == 'Cuisine':
        return (47, 84, 150)
    elif style == 'Dish':
        return (255,0,0)
    elif "Group" in style:
        return (196,89,17)
    else:
        return (0,0,0)

def change_document(document):
    pass

def write_paragraphs_to_file(path):
    document = docx.Document(path)
    styles = document.styles
    print([style for style in styles])
    for paragraph in document.paragraphs:
        if requires_update(paragraph):
            try:
                new_style = styles[determine_style(paragraph)].name
            except:
                new_style = 'Normal'
            paragraph.style = styles[new_style]
            rgb = determine_font_RGB_color(new_style)
            for run in paragraph.runs:
                run.font.color.rgb = docx.shared.RGBColor(rgb[0], rgb[1], rgb[2])
    document.save('StyledIngredients.docx')








path = "../../../IngredientPairs.docx"
print(os.path.abspath(path))

doc = docx.Document(path)
write_paragraphs_to_file(path)